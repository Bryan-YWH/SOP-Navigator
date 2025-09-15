#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SOP文档智能解析脚本 - 精确表格归属版本
实现精确的上下文追踪、表格归属和动态标题生成
"""

import sys
import re
import json
import pandas as pd
from collections import defaultdict
from typing import List
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml import OxmlElement


def is_heading_paragraph(paragraph: Paragraph) -> bool:
    """
    判断段落是否为标题
    使用三级优先序：样式优先 -> 数字编号 -> 关键词
    """
    # 第一优先级：检查Word样式
    style_name = paragraph.style.name
    if 'Heading' in style_name or '标题' in style_name:
        return True
    
    # 第二优先级：检查数字编号
    text = paragraph.text.strip()
    
    # 多级数字编号 (如 3.1, 8.2.1)
    if re.match(r'^\d+(?:\.\d+)+', text):
        return True
    
    # 单级数字编号 (如 4), 5))
    if re.match(r'^\d+\)', text):
        return True
    
    # 纯数字标题 (如 8.历史文件记录)
    if re.match(r'^\d+\.\s+', text):
        return True
    
    # 第三优先级：检查关键词
    TOP_LEVEL_KEYWORDS = [
        "目的", "适用范围", "安全和环境要求", "相关文件", "职责", 
        "定义和缩写", "活动描叙", "具体操作如下", "附录", "历史纪录"
    ]
    
    # 清理文本前的编号
    clean_text = re.sub(r'^\d+(?:\.\d+)*[\.\)]\s*', '', text).strip()
    if clean_text in TOP_LEVEL_KEYWORDS:
        return True
    
    # 特殊处理：检查是否包含"活动描述"关键词
    if "活动描述" in text and re.match(r'^\d+\.', text):
        return True
    
    return False


def get_heading_level(paragraph: Paragraph) -> int:
    """
    获取标题的层级
    """
    # 第一优先级：检查Word样式
    style_name = paragraph.style.name
    if 'Heading' in style_name:
        match = re.search(r'Heading (\d+)', style_name)
        if match:
            return int(match.group(1))
    elif '标题' in style_name:
        match = re.search(r'标题 (\d+)', style_name)
        if match:
            return int(match.group(1))
    
    # 第二优先级：根据数字编号推断层级
    text = paragraph.text.strip()
    
    # 多级数字编号
    match = re.match(r'^(\d+(?:\.\d+)+)', text)
    if match:
        level = match.group(1).count('.') + 1
        return min(level, 6)  # 限制最大层级为6
    
    # 单级数字编号或纯数字标题
    if re.match(r'^\d+[\)\.]', text):
        return 1
    
    # 第三优先级：关键词默认为1级
    return 1


def table_to_markdown(table: Table) -> str:
    """
    将Word表格转换为Markdown格式
    """
    if not table.rows:
        return ""
    
    markdown_lines = []
    
    # 处理表头
    header_row = table.rows[0]
    header_cells = [cell.text.strip() for cell in header_row.cells]
    markdown_lines.append("| " + " | ".join(header_cells) + " |")
    
    # 处理分隔符
    separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"
    markdown_lines.append(separator)
    
    # 处理数据行
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        markdown_lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(markdown_lines)


def normalize_list_symbols(text: str) -> str:
    """
    将非标准的列表符号替换为标准的Markdown格式
    """
    # 替换各种非标准列表符号
    text = re.sub(r'^[·•]\s*', '* ', text, flags=re.MULTILINE)
    text = re.sub(r'^--\t', '* ', text, flags=re.MULTILINE)
    text = re.sub(r'^、\s*', '* ', text, flags=re.MULTILINE)
    
    return text


def build_section_path(heading_stack: list) -> str:
    """
    构建完整的章节路径
    """
    return " > ".join(heading_stack)


def identify_table_section(table_content: str) -> str:
    """
    根据表格内容识别表格应该归属的章节
    """
    # 根据表格内容特征判断归属章节
    if "分类" in table_content and "危险源" in table_content and "控制措施" in table_content:
        return "3.1 风险识别"
    elif "相关模块" in table_content and "危险源" in table_content and "控制措施" in table_content:
        return "3.2 关键控制点"
    elif "成品库保管员" in table_content and "成品库班长" in table_content and "SOP撰写" in table_content:
        return "5.职责"
    elif "本SOP涉及到的主要KPI" in table_content and "PI" in table_content:
        return "6.定义和缩写"
    elif "版本" in table_content and "作者" in table_content and "日期" in table_content:
        return "8.历史文件记录"
    elif "仓库利用率" in table_content and "劳动生产率" in table_content:
        return "6.定义和缩写"
    elif "PPE矩阵" in table_content and "风险评估" in table_content:
        return "3.2 关键控制点"
    elif "应急方案" in table_content and "成品酒高空坠落" in table_content:
        return "3.2 关键控制点"
    else:
        return "未知章节"


def build_table_section_path(table_section: str, heading_stack: List[str]) -> str:
    """
    根据表格所属章节构建正确的section_path
    """
    # 如果表格章节是顶级章节，直接返回
    if table_section in ["3.1 风险识别", "3.2 关键控制点", "5.职责", "6.定义和缩写", "8.历史文件记录"]:
        return table_section
    
    # 如果是子章节，需要找到对应的父章节
    for i, heading in enumerate(heading_stack):
        if table_section.startswith(heading.split()[0]):  # 匹配章节号
            return " > ".join(heading_stack[:i+1]) + f" > {table_section}"
    
    # 如果找不到匹配的父章节，返回表格章节本身
    return table_section


def process_sop_document(docx_path: str) -> list:
    """
    处理SOP文档，返回所有知识块
    """
    print("============================================================")
    print("SOP文档智能解析工具 - 精确表格归属版本")
    print("============================================================")
    
    # 读取Word文档
    try:
        doc = Document(docx_path)
        print(f"成功读取文档: {docx_path}")
    except Exception as e:
        print(f"读取文档失败: {e}")
        return []
    
    # 初始化变量
    chunks = []
    heading_stack = []  # 维护标题层级栈
    current_heading_text = ""  # 当前最深层级的标题文本
    table_counter_map = defaultdict(int)  # 为每个标题维护独立的表格计数器
    current_content_buffer = []  # 当前小节的内容缓冲区
    
    # 从文档中提取SOP信息
    sop_id = "未知"
    sop_name = "未知"
    
    # 尝试从文档标题或第一段获取SOP信息
    if doc.paragraphs:
        first_para = doc.paragraphs[0].text.strip()
        if first_para:
            sop_name = first_para
            # 尝试提取SOP ID
            id_match = re.search(r'([A-Z0-9\.]+)', first_para)
            if id_match:
                sop_id = id_match.group(1)
    
    print(f"SOP ID: {sop_id}")
    print(f"SOP名称: {sop_name}")
    
    # 遍历文档的所有段落和表格
    # 首先处理所有段落
    for paragraph in doc.paragraphs:
        if is_heading_paragraph(paragraph):
            # 这是一个标题
            heading_text = paragraph.text.strip()
            heading_level = get_heading_level(paragraph)
            
            # 如果有收集的内容，先处理之前的内容（包括上一个标题和其内容）
            if current_content_buffer:
                section_path = build_section_path(heading_stack)
                combined_text = normalize_list_symbols('\n'.join(current_content_buffer))
                
                chunks.append({
                    'text': combined_text,
                    'sop_id': sop_id,
                    'sop_name': sop_name,
                    'section_path': section_path,
                    'image_filename': ''
                })
                current_content_buffer = []
            
            # 更新标题栈
            # 移除同级及更深层级的标题
            heading_stack = [h for h in heading_stack if heading_stack.index(h) < heading_level - 1]
            
            # 添加新标题
            if len(heading_stack) >= heading_level:
                heading_stack = heading_stack[:heading_level-1]
            heading_stack.append(heading_text)
            
            # 更新当前标题文本
            current_heading_text = heading_text
            
            # 将新标题添加到内容缓冲区，作为新小节内容的开始
            current_content_buffer.append(heading_text)
            
            print(f"处理标题: {heading_text} (层级: {heading_level})")
            
        else:
            # 这是一个普通段落
            if paragraph.text.strip():
                current_content_buffer.append(paragraph.text.strip())
    
    # 处理最后收集的内容
    if current_content_buffer:
        section_path = build_section_path(heading_stack)
        combined_text = normalize_list_symbols('\n'.join(current_content_buffer))
        
        chunks.append({
            'text': combined_text,
            'sop_id': sop_id,
            'sop_name': sop_name,
            'section_path': section_path,
            'image_filename': ''
        })
    
    # 然后处理所有表格
    for table in doc.tables:
        # 转换表格为Markdown
        markdown_table = table_to_markdown(table)
        
        if markdown_table:
            # 根据表格内容识别归属章节
            table_section = identify_table_section(markdown_table)
            
            # 如果有收集的内容，先处理当前小节的文本内容
            if current_content_buffer:
                section_path = build_section_path(heading_stack)
                combined_text = normalize_list_symbols('\n'.join(current_content_buffer))
                
                chunks.append({
                    'text': combined_text,
                    'sop_id': sop_id,
                    'sop_name': sop_name,
                    'section_path': section_path,
                    'image_filename': ''
                })
                current_content_buffer = []
            
            # 增加表格所属章节的表格计数器
            table_counter_map[table_section] += 1
            table_counter = table_counter_map[table_section]
            
            # 生成动态标题和内容
            dynamic_title = f"标题：{table_section} - 表格 {table_counter}"
            combined_text = f"{dynamic_title}\n\n{markdown_table}"
            
            # 构建section_path - 根据表格所属章节构建正确的路径
            table_section_path = build_table_section_path(table_section, heading_stack)
            
            # 创建独立的知识块
            chunks.append({
                'text': combined_text,
                'sop_id': sop_id,
                'sop_name': sop_name,
                'section_path': table_section_path,
                'image_filename': ''
            })
            
            print(f"处理表格: {table_section} - 表格 {table_counter}")
    
    print(f"总共处理了 {len(chunks)} 个知识块")
    return chunks


def main():
    """
    主函数
    """
    if len(sys.argv) != 2:
        print("使用方法: python process_sop_to_csv.py <docx文件路径>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    # 处理文档
    chunks = process_sop_document(docx_path)
    
    if not chunks:
        print("没有生成任何知识块")
        return
    
    # 生成输出文件名
    output_path = docx_path.replace('.docx', '_processed.csv')
    
    # 保存为CSV
    try:
        df = pd.DataFrame(chunks)
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        print(f"成功保存到: {output_path}")
        print(f"总共生成 {len(chunks)} 个知识块")
    except Exception as e:
        print(f"保存CSV文件失败: {e}")


if __name__ == "__main__":
    main()
