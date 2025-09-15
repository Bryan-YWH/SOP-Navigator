#!/usr/bin/env python3
"""
将 .docx 文档解析为显式的嵌套 JSON，支持健壮的混合识别逻辑。

功能要点：
- 使用 sys.argv 接收输入文件路径并进行基本检查。
- 使用 python-docx 读取段落与样式，实现三级优先序的混合识别逻辑。
- 普通段落归属到其上方最近的标题的 content 列表中。
- 自动识别并处理文档中的表格，将表格内容转换为Markdown格式。
- 输出 JSON: { sop_id, sop_name, sections }，sections 为顶层标题列表，
  每个标题包含 { title, level, content, subsections }。

三级优先序混合识别逻辑：
1. 第一优先级 - 检查样式：paragraph.style.name 匹配 'Heading X' 或 '标题 X'
2. 第二优先级 - 检查数字编号：使用正则表达式匹配数字编号格式推断标题级别
   - 规则A (多级): 匹配如 3.1, 8.2.1 这样的格式，数字和文字间可能没有空格
   - 规则B (单级): 匹配如 4), 5), 10) 这样的格式，统一视为一级标题
3. 第三优先级 - 检查关键词：检查纯文本内容是否完全等于预定义的关键词列表

注意：
- 文档标题(样式名 'Title' 或中文 '标题')不会作为章节加入，而是作为 sop_name 的优先来源。
- 若无法检测到文档标题，则 sop_name 取文档第一段非空文本；若仍不可得，则回退为 sop_id。
"""

from __future__ import annotations

import json
import os
import re
import sys
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document  # type: ignore
    from docx.oxml.table import CT_Tbl  # type: ignore
    from docx.oxml.text.paragraph import CT_P  # type: ignore
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
except ImportError:  # pragma: no cover
    sys.stderr.write(
        "[ERROR] 未找到 python-docx 库，请先安装：pip install python-docx\n"
    )
    sys.exit(1)

# 预定义的关键词列表，用于识别纯文本标题
TOP_LEVEL_KEYWORDS = [
    "目的", "适用范围", "安全和环境要求", "相关文件", "职责", 
    "定义和缩写", "活动描叙", "具体操作如下", "附录", "历史纪录"
]

HeadingNode = Dict[str, Any]


def extract_heading_level_from_style(style_name: Optional[str]) -> Optional[int]:
    """根据段落样式名推断标题级别（第一优先级）。

    支持示例：
    - 'Heading 1', 'Heading 2', ..., 'Heading 10'
    - '标题 1', '标题 2', ..., '标题 10'（中文界面）
    返回对应的整数级别（1-10）；若非标题样式则返回 None。
    """
    if not style_name:
        return None

    # 统一去除两端空白，便于匹配
    normalized = style_name.strip()

    # 英文样式：Heading 1/2/3 ...
    m = re.match(r"^Heading\s+([1-9][0-9]*)$", normalized, re.IGNORECASE)
    if m:
        level = int(m.group(1))
        return level if 1 <= level <= 10 else None

    # 中文样式：标题 1/2/3 ...
    m = re.match(r"^标题\s*([1-9][0-9]*)$", normalized)
    if m:
        level = int(m.group(1))
        return level if 1 <= level <= 10 else None

    return None


def extract_heading_level_from_text(text: str) -> Optional[int]:
    """根据段落文本内容推断标题级别（第二优先级）。

    规则A (多级): 匹配如 3.1, 8.2.1 这样的格式，数字和文字间可能没有空格
    规则B (单级): 匹配如 4), 5), 10) 这样的格式，统一视为一级标题
    
    返回对应的整数级别；若文本不符合编号格式则返回 None。
    """
    if not text:
        return None
    
    text = text.strip()
    
    # 规则A: 多级数字编号 (如 3.1, 8.2.1, 2.1.1.1 等)
    # 使用 \s* 而不是 \s+ 来兼容数字和文字间没有空格的情况
    multi_level_match = re.match(r'^(\d+(?:\.\d+)+)\s*', text)
    if multi_level_match:
        # 计算点号数量来确定级别
        number_part = multi_level_match.group(1)
        level = number_part.count('.') + 1
        return level if 1 <= level <= 10 else None
    
    # 规则B: 单级括号编号 (如 4), 5), 10) 等)
    # 统一视为一级标题
    single_level_match = re.match(r'^\d+\)\s*', text)
    if single_level_match:
        return 1
    
    return None


def extract_heading_level_from_keywords(text: str) -> Optional[int]:
    """根据关键词列表推断标题级别（第三优先级）。

    检查清理后的文本是否完全等于预定义的关键词列表中的某一项。
    如果是，则视为一级标题。
    """
    if not text:
        return None
    
    # 清理文本：移除可能的前置编号和特殊字符
    clean_text = re.sub(r'^[\d\.\)\s]+', '', text).strip()
    
    # 检查是否完全匹配关键词列表
    if clean_text in TOP_LEVEL_KEYWORDS:
        return 1
    
    return None


def extract_heading_level(style_name: Optional[str], text: str) -> Optional[int]:
    """三级优先序混合识别逻辑。
    
    第一优先级：检查Word样式
    第二优先级：检查数字编号（多级和单级）
    第三优先级：检查关键词列表
    
    返回标题级别（1-10），若非标题则返回 None。
    """
    # 第一优先级：检查Word样式
    level = extract_heading_level_from_style(style_name)
    if level is not None:
        return level
    
    # 第二优先级：检查数字编号
    level = extract_heading_level_from_text(text)
    if level is not None:
        return level
    
    # 第三优先级：检查关键词列表
    level = extract_heading_level_from_keywords(text)
    if level is not None:
        return level
    
    return None


def is_document_title_style(style_name: Optional[str]) -> bool:
    """判断样式是否为文档标题(非 Heading 层级)。

    常见样式名：'Title' 或中文界面中的 '标题'（注意与 '标题 1' 区分）。
    """
    if not style_name:
        return False
    name = style_name.strip()
    if name.lower() == "title":
        return True
    # 精确等于"标题"视为文档标题，避免与"标题 1/2..."混淆
    if name == "标题":
        return True
    return False


def convert_table_to_markdown(table: Table) -> str:
    """将python-docx的Table对象转换为Markdown表格格式的字符串。
    
    参数:
        table: python-docx的Table对象
        
    返回:
        Markdown格式的表格字符串
    """
    if not table or not table.rows:
        return ""

    # 构建Markdown表格的字符串
    markdown_string = ""
    
    # 获取表格数据
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # 获取单元格文本，去除多余空白
            cell_text = cell.text.strip()
            # 转义Markdown特殊字符
            cell_text = cell_text.replace('|', '\\|')
            row_data.append(cell_text)
        table_data.append(row_data)
    
    if not table_data:
        return ""

    # 1. 处理表头
    header = "| " + " | ".join(table_data[0]) + " |"
    markdown_string += header + "\n"

    # 2. 处理分隔符行
    separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
    markdown_string += separator + "\n"

    # 3. 处理表格内容行
    for row in table_data[1:]:
        # 确保行数据长度与表头一致
        while len(row) < len(table_data[0]):
            row.append("")
        content_row = "| " + " | ".join(row[:len(table_data[0])]) + " |"
        markdown_string += content_row + "\n"

    return markdown_string


def make_node(title: str, level: int) -> HeadingNode:
    """创建一个标题节点，将标题同时作为第一条内容。"""
    return {
        "title": title,
        "level": level,
        "content": [title],  # 将标题作为第一条内容
        "images": [],  # type: List[str]
        "subsections": [],  # type: List[HeadingNode]
    }


def docx_to_nested_json(input_path: str) -> Dict[str, Any]:
    """解析给定 .docx 文档，返回目标 JSON 结构的 Python 字典。"""
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"文件不存在: {input_path}")
    if not input_path.lower().endswith(".docx"):
        raise ValueError("输入文件必须为 .docx 格式")

    sop_id = os.path.splitext(os.path.basename(input_path))[0]
    
    try:
        # 使用 python-docx 解析文档
        doc = Document(input_path)
    except Exception as exc:  # pragma: no cover
        raise ValueError(f"无法读取 Word 文档: {exc}")

    # 提取文档标题(sop_name)
    sop_name: Optional[str] = None
    
    # 首先尝试从文档标题样式获取
    for paragraph in doc.paragraphs:
        if is_document_title_style(paragraph.style.name) and paragraph.text.strip():
            sop_name = paragraph.text.strip()
            break
    
    # 如果没有找到文档标题样式，则从第一段非空文本获取
    if not sop_name:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and not re.match(r'^\d+\.', text):  # 不是数字编号开头
                sop_name = text
                break
    
    # 如果仍然没有找到，使用sop_id作为sop_name
    if not sop_name:
        sop_name = sop_id

    sections: List[HeadingNode] = []
    # 栈保存 (level, node)
    heading_stack: List[Tuple[int, HeadingNode]] = []

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue

        # 检查是否为标题
        level = extract_heading_level(paragraph.style.name, text)

        if level is not None:
            # 新的标题节点。维护单调栈，使其父子关系正确。
            new_node = make_node(text, level)

            # 弹出栈中当前级别及更深的标题，直到找到父级(level 更小)
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()

            if not heading_stack:
                # 顶级标题
                sections.append(new_node)
            else:
                # 作为最近上级标题的子节点
                parent_node = heading_stack[-1][1]
                parent_node["subsections"].append(new_node)

            # 将当前标题压栈
            heading_stack.append((level, new_node))
        else:
            # 普通段落
            if heading_stack:
                heading_stack[-1][1]["content"].append(text)
            else:
                # 若在任何标题出现之前出现正文，则不纳入 sections。
                # 可根据业务需要改为挂入一个"前言"节点。
                pass

    # 处理文档中的所有表格
    for table in doc.tables:
        if table.rows:  # 确保表格不为空
            # 将表格转换为Markdown格式
            markdown_table = convert_table_to_markdown(table)
            if markdown_table and heading_stack:
                # 表格归属到最近的标题
                heading_stack[-1][1]["content"].append(markdown_table)

    return {
        "sop_id": sop_id,
        "sop_name": sop_name,
        "sections": sections,
    }


def main(argv: List[str]) -> int:
    if len(argv) < 2:
        sys.stderr.write(
            "用法: python word_to_json.py <输入文件.docx>\n"
        )
        return 1

    input_path = argv[1]
    try:
        data = docx_to_nested_json(input_path)
    except (FileNotFoundError, ValueError) as exc:
        sys.stderr.write(f"[ERROR] {exc}\n")
        return 1
    except Exception as exc:  # pragma: no cover
        # 捕获未知错误，便于定位问题
        sys.stderr.write(f"[ERROR] 未预期的异常: {exc}\n")
        return 1

    # 输出至同名 .json
    output_path = os.path.splitext(input_path)[0] + ".json"
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as exc:  # pragma: no cover
        sys.stderr.write(f"[ERROR] 写入 JSON 失败: {exc}\n")
        return 1

    print(f"已生成: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))